# backend/app.py
import os
import ollama
import json
import uuid
from flask import Flask, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from datetime import datetime
import logging

# Imports pour l'extraction de texte
import PyPDF2
import docx # python-docx
import openpyxl

# Configuration du logging
logging.basicConfig(level=logging.INFO)

# Configuration de Flask
# Le static_folder pointe vers le dossier frontend relatif à ce fichier app.py
app = Flask(__name__, static_folder='../frontend', static_url_path='')

# Configuration des dossiers
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
LOG_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
# Extensions autorisées pour l'upload ET l'extraction (ajustées pour les libs utilisées)
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx', 'txt'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(LOG_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['LOG_FOLDER'] = LOG_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 # Limite d'upload à 16MB

# --- État Global ---
# NOUVEAU: Modèle par défaut mis à jour selon ta demande
current_model = 'deepseek-r1:14b'
current_conversation_history = []
current_conversation_id = None
# Dictionnaire pour stocker le contexte des fichiers par conversation
# Format: { "conversation_id": {"filename": "...", "content": "..."} }
uploaded_file_context = {}

# --- Fonctions Utilitaires ---
def allowed_file(filename):
    """Vérifie si l'extension du fichier est autorisée."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_log_filename(conv_id):
    """Génère le chemin complet vers le fichier log pour un ID de conversation."""
    return os.path.join(app.config['LOG_FOLDER'], f"conversation_{conv_id}.json")

def save_log(conv_id, history):
    """Sauvegarde l'historique d'une conversation dans un fichier JSON."""
    if not conv_id:
        logging.warning("Tentative de sauvegarde sans ID de conversation.")
        return
    log_file = generate_log_filename(conv_id)
    try:
        with open(log_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=4)
        logging.info(f"Conversation {conv_id} sauvegardée dans {log_file}")
    except Exception as e:
        logging.error(f"Erreur lors de la sauvegarde du log {conv_id}: {e}")

def load_log(conv_id):
    """Charge l'historique d'une conversation depuis un fichier JSON."""
    log_file = generate_log_filename(conv_id)
    if os.path.exists(log_file):
        try:
            with open(log_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logging.error(f"Erreur lors du chargement du log {conv_id}: {e}")
            return []
    return []

def extract_text_from_file(filepath):
    """Extrait le texte d'un fichier PDF, DOCX, XLSX ou TXT."""
    filename = os.path.basename(filepath)
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    text = ""
    try:
        if ext == 'pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                # Vérifier si le PDF est chiffré (simple vérification)
                if reader.is_encrypted:
                     logging.warning(f"Le fichier PDF '{filename}' est chiffré et ne peut pas être lu.")
                     # Tenter de déchiffrer avec un mot de passe vide (pour certains cas)
                     try:
                         reader.decrypt('')
                     except Exception as decrypt_error:
                          logging.error(f"Échec du déchiffrement (sans mdp) de {filename}: {decrypt_error}")
                          return None # Échec si chiffré et déchiffrement échoue
                # Extraire le texte page par page
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif ext == 'docx':
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Optionnel: Ajouter le texte des tableaux
            # for table in doc.tables:
            #     for row in table.rows:
            #         for cell in row.cells:
            #             text += cell.text + "\t"
            #         text += "\n"
        elif ext == 'xlsx':
            workbook = openpyxl.load_workbook(filepath, data_only=True, read_only=True) # data_only=True pour avoir les valeurs des formules
            for sheetname in workbook.sheetnames:
                sheet = workbook[sheetname]
                for row in sheet.iter_rows():
                    row_text = []
                    for cell in row:
                        if cell.value is not None:
                            # Formater les nombres et dates de manière plus lisible si nécessaire
                            # Ici, on convertit simplement en chaîne
                            row_text.append(str(cell.value).strip())
                    if row_text:
                        text += "\t".join(row_text) + "\n"
        elif ext == 'txt':
            # Essayer plusieurs encodages courants si l'UTF-8 échoue
            encodings_to_try = ['utf-8', 'latin-1', 'windows-1252']
            for enc in encodings_to_try:
                try:
                    with open(filepath, 'r', encoding=enc) as f:
                        text = f.read()
                    logging.info(f"Fichier texte lu avec l'encodage : {enc}")
                    break # Sortir de la boucle si la lecture réussit
                except UnicodeDecodeError:
                    logging.warning(f"Échec de la lecture de {filename} avec {enc}, essai suivant...")
                except Exception as e_inner: # Capturer d'autres erreurs de lecture potentielles
                     logging.error(f"Erreur inattendue lors de la lecture de {filename} avec {enc}: {e_inner}")
                     text = None # Marquer comme échec
                     break
            if text is None: # Si tous les encodages ont échoué
                 logging.error(f"Impossible de décoder le fichier texte {filename} avec les encodages testés.")
                 return None

        else:
            logging.warning(f"Type de fichier non supporté pour l'extraction: {ext}")
            return None

        # Vérifier si du texte a été extrait
        extracted_text = text.strip()
        if not extracted_text:
            logging.warning(f"Aucun texte n'a pu être extrait de {filename} (fichier vide ou format non reconnu?).")
            return None

        logging.info(f"Texte extrait de {filename} (longueur: {len(extracted_text)} caractères)")
        return extracted_text

    except FileNotFoundError:
        logging.error(f"Le fichier spécifié n'a pas été trouvé : {filepath}")
        return None
    except PyPDF2.errors.PdfReadError as pdf_error:
         logging.error(f"Erreur de lecture PDF pour {filename}: {pdf_error}. Fichier corrompu ou format non supporté?")
         return None
    except Exception as e:
        logging.error(f"Erreur générale lors de l'extraction du texte de {filename} ({ext}): {e}", exc_info=True) # exc_info=True pour la traceback
        return None

# --- Routes de l'API ---

@app.route('/')
def index():
    """Sert la page HTML principale."""
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/set_model', methods=['POST'])
def set_model():
    """Change le modèle Ollama utilisé et démarre une nouvelle conversation."""
    global current_model, uploaded_file_context
    data = request.json
    model_size = data.get('model_size') # C'est la 'key' comme "7b", "r1-14b" etc.

    if not model_size:
         return jsonify({"success": False, "message": "Taille du modèle non fournie."}), 400

    # *** NOUVEAU: Mapping mis à jour avec tes noms de modèles exacts ***
    model_mapping = {
        "1.5b": "deepseek-r1:1.5b",
        "7b": "deepseek-r1:7b",
        "8b": "deepseek-r1:8b",      # Correction du typo "dedeepseek" effectuée ici
        "14b": "deepseek-r1:14b",     # Mappe la valeur "14b" du dropdown
        "32b": "deepseek-r1:32b",
        "70b": "deepseek-r1:70b",
        "671b": "deepseek-r1:671b",    # Attention à la RAM nécessaire pour ce modèle !
        "r1-14b": "deepseek-r1:14b"   # Mappe la valeur "r1-14b" du dropdown (si elle existe)
        # Assure-toi que les 'value' dans ton HTML correspondent aux clés ici
    }

    selected_model_name = model_mapping.get(model_size)

    if not selected_model_name:
         logging.error(f"Aucun modèle correspondant trouvé dans le mapping pour la clé: {model_size}")
         return jsonify({"success": False, "message": f"Clé de modèle '{model_size}' inconnue."}), 400

    logging.info(f"Tentative de changement de modèle vers : {selected_model_name} (demandé via '{model_size}')")

    try:
        # Vérifie si le modèle existe localement via l'API ollama
        ollama.show(selected_model_name)
        # Si ça ne lève pas d'exception, le modèle existe
        current_model = selected_model_name
        logging.info(f"Modèle changé avec succès pour : {current_model}")
        # Démarrer une nouvelle conversation ET effacer le contexte fichier associé
        start_new_chat_internal() # Ceci gère la réinitialisation
        return jsonify({
            "success": True,
            "message": f"Modèle changé pour {current_model}. Nouvelle conversation démarrée.",
            "conversation_id": current_conversation_id
        })
    except ollama.ResponseError as e:
         logging.error(f"Erreur lors de la vérification/changement de modèle pour '{selected_model_name}': Status {e.status_code}, Erreur: {e.error}. Le modèle n'existe peut-être pas localement ou n'est pas accessible.")
         # Ne pas changer current_model si l'autre n'est pas valide
         return jsonify({
             "success": False,
             "message": f"Modèle '{selected_model_name}' non trouvé ou invalide. Vérifiez les modèles disponibles via 'ollama list'."}), 400 # 400 Bad Request est approprié
    except Exception as e:
         logging.error(f"Erreur inattendue lors du changement de modèle: {e}", exc_info=True)
         return jsonify({"success": False, "message": "Erreur interne lors du changement de modèle."}), 500


@app.route('/chat', methods=['POST'])
def chat():
    """Gère l'envoi d'un message utilisateur et la réception de la réponse de l'IA."""
    global current_conversation_history, current_conversation_id, uploaded_file_context
    data = request.json
    user_input = data.get('message')

    if not user_input:
        return jsonify({"error": "Message vide reçu."}), 400

    # S'assure qu'une conversation est active
    if not current_conversation_id:
        start_new_chat_internal()

    # Ajoute le message utilisateur à l'historique PERSISTANT
    current_conversation_history.append({'role': 'user', 'content': user_input})

    # Crée une copie de l'historique pour l'envoyer à Ollama
    history_for_ollama = list(current_conversation_history)

    # Injecte le contexte du fichier si disponible pour cette conversation
    file_ctx = uploaded_file_context.get(current_conversation_id)
    if file_ctx and file_ctx.get('content'): # Vérifie aussi que le contenu existe
        logging.info(f"Injection du contexte du fichier '{file_ctx['filename']}' ({len(file_ctx['content'])} chars) pour la conversation {current_conversation_id}")

        # Message système expliquant le contexte fichier
        # NOTE: Peut être très long!
        context_prompt = (
            f"CONTEXTE IMPORTANT FOURNI PAR L'UTILISATEUR :\n"
            f"Le contenu suivant provient du fichier '{file_ctx['filename']}'. "
            f"Utilise ces informations pour répondre à la question de l'utilisateur si elle s'y rapporte.\n"
            f"--- DEBUT DU CONTENU DU FICHIER ---\n"
            f"{file_ctx['content']}\n"
            f"--- FIN DU CONTENU DU FICHIER ---\n"
            f"FIN DU CONTEXTE."
            # Déplacé la partie "Réponds maintenant..." à la fin du message user pour clarté
        )
        context_message = {"role": "system", "content": context_prompt}

        # Insérer le contexte système au début de la copie de l'historique
        # C'est généralement la meilleure place pour les instructions système
        history_for_ollama.insert(0, context_message)

        # Optionnel: Ajouter une note à la fin du dernier message utilisateur pour rappeler le contexte
        # history_for_ollama[-1]['content'] += "\n(Rappel: J'ai fourni un fichier nommé '" + file_ctx['filename'] + "' comme contexte.)"

        # Log d'avertissement si le contexte est très volumineux
        if len(file_ctx['content']) > 15000: # Seuil (en caractères, pas tokens)
             logging.warning(f"Le contexte injecté fait {len(file_ctx['content'])} caractères. Risque élevé de dépasser la limite de contexte du modèle '{current_model}'.")

    else:
        logging.info("Aucun contexte fichier à injecter pour cette requête.")


    logging.info(f"Envoi à Ollama ({current_model}) avec historique de {len(history_for_ollama)} messages.")
    # Pour le debug, ne pas logger l'historique complet s'il contient des données sensibles/longues
    # logging.debug(f"Historique envoyé (avec contexte fichier éventuel): {history_for_ollama}")

    try:
        # Appel à l'API Ollama
        response = ollama.chat(
            model=current_model,
            messages=history_for_ollama, # Utilise l'historique potentiellement augmenté
            stream=False # Garde False pour une réponse complète
        )

        ai_response = response['message']['content']

        # Ajoute UNIQUEMENT la réponse de l'IA à l'historique PERSISTANT
        current_conversation_history.append({'role': 'assistant', 'content': ai_response})

        # Sauvegarde l'historique PERSISTANT (sans le contexte fichier systémique)
        save_log(current_conversation_id, current_conversation_history)

        # Renvoie la réponse au frontend
        return jsonify({"response": ai_response})

    except ollama.ResponseError as e:
        logging.error(f"Erreur Ollama: Status {e.status_code}, Erreur: {e.error}")
        # Retirer le dernier message utilisateur de l'historique réel si l'appel échoue
        if current_conversation_history and current_conversation_history[-1]['role'] == 'user':
             current_conversation_history.pop()
        # Construire un message d'erreur plus utile
        error_message = f"Erreur lors de la communication avec Ollama ({e.status_code})."
        if "context window" in str(e.error).lower() or "maximum sequence length" in str(e.error).lower():
             error_message += " Le contexte (incluant potentiellement le fichier) est probablement trop long pour ce modèle. Essayez avec un fichier plus petit ou une conversation plus courte."
        elif "not found" in str(e.error).lower():
             error_message += f" Le modèle '{current_model}' semble ne pas être disponible."
        else:
            error_message += f" Détail: {e.error}"
        return jsonify({"error": error_message}), 500 # 500 Internal Server Error ou 502 Bad Gateway
    except Exception as e:
        logging.error(f"Erreur inattendue lors du chat: {e}", exc_info=True)
        # Retirer aussi le dernier message user en cas d'erreur générique
        if current_conversation_history and current_conversation_history[-1]['role'] == 'user':
             current_conversation_history.pop()
        return jsonify({"error": "Une erreur interne inattendue est survenue lors du traitement de votre message."}), 500


@app.route('/upload', methods=['POST'])
def upload_file():
    """Gère l'upload d'un fichier, extrait son texte et le stocke."""
    global uploaded_file_context, current_conversation_id

    if 'file' not in request.files:
        return jsonify({"error": "Aucun fichier sélectionné"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nom de fichier vide"}), 400

    if not allowed_file(file.filename):
        logging.warning(f"Tentative d'upload d'un type de fichier non autorisé: {file.filename}")
        allowed_types_str = ", ".join(ALLOWED_EXTENSIONS)
        return jsonify({"error": f"Type de fichier non autorisé. Types permis : {allowed_types_str}."}), 400

    filename = secure_filename(file.filename)
    # Optionnel : ajouter un timestamp ou UUID pour éviter les collisions de noms
    # unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
    # filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    try:
        file.save(filepath)
        logging.info(f"Fichier '{filename}' sauvegardé dans {filepath}")

        # Extraire le texte du fichier sauvegardé
        extracted_text = extract_text_from_file(filepath)

        if extracted_text is not None:
             # S'assurer qu'une conversation est active
             if not current_conversation_id:
                 start_new_chat_internal() # Démarre une nouvelle conversation si aucune n'est active

             # Stocker le texte extrait associé à l'ID de la conversation courante
             # Écrase le contexte précédent pour cette conversation s'il existait
             uploaded_file_context[current_conversation_id] = {
                 "filename": filename,
                 "content": extracted_text
             }
             logging.info(f"Texte de '{filename}' (environ {len(extracted_text)} chars) stocké pour la conversation {current_conversation_id}.")
             return jsonify({
                 "success": True,
                 "message": f"Fichier '{filename}' uploadé et traité ({len(extracted_text)} caractères extraits). Vous pouvez maintenant poser des questions sur son contenu.",
                 "filename": filename
             })
        else:
             # L'extraction a échoué (erreur déjà logguée dans extract_text_from_file)
             # Garder le fichier uploadé ? Ou le supprimer ? Pour l'instant on le garde.
             return jsonify({"error": f"Fichier '{filename}' uploadé mais impossible d'en extraire le contenu. Vérifiez le format ou si le fichier est protégé/corrompu."}), 422 # 422 Unprocessable Entity

    except Exception as e:
        # Erreur probable lors de file.save() ou autre imprévu
        logging.error(f"Erreur lors de la sauvegarde ou traitement du fichier '{filename}': {e}", exc_info=True)
        return jsonify({"error": "Erreur serveur lors de la sauvegarde ou du traitement initial du fichier."}), 500


def start_new_chat_internal():
    """Fonction interne sécurisée pour démarrer un nouveau chat."""
    global current_conversation_history, current_conversation_id, uploaded_file_context
    # Nettoie le contexte fichier de la conversation précédente (si elle existait)
    if current_conversation_id in uploaded_file_context:
        logging.info(f"Nettoyage du contexte fichier pour l'ancienne conversation {current_conversation_id}")
        del uploaded_file_context[current_conversation_id]

    # Génère un nouvel ID unique pour la nouvelle conversation
    current_conversation_id = str(uuid.uuid4())
    # Réinitialise l'historique en mémoire
    current_conversation_history = []
    # Pas besoin de toucher uploaded_file_context pour le nouvel ID, il est vide par défaut
    logging.info(f"Nouvelle conversation démarrée avec ID: {current_conversation_id}")


@app.route('/new_chat', methods=['POST'])
def new_chat():
    """Route API pour démarrer une nouvelle conversation."""
    start_new_chat_internal() # Appelle la fonction interne qui fait le travail
    return jsonify({"success": True, "message": "Nouveau chat démarré.", "conversation_id": current_conversation_id})


@app.route('/history', methods=['GET'])
def get_history_list():
    """Retourne la liste des conversations sauvegardées."""
    try:
        log_files = [f for f in os.listdir(app.config['LOG_FOLDER']) if f.startswith("conversation_") and f.endswith(".json")]
        conversations = []
        # Trier par date de modification (plus récent en premier)
        log_files.sort(key=lambda f: os.path.getmtime(os.path.join(app.config['LOG_FOLDER'], f)), reverse=True)

        for filename in log_files:
            conv_id = filename.replace("conversation_", "").replace(".json", "")
            try:
                filepath = os.path.join(app.config['LOG_FOLDER'], filename)
                mtime = os.path.getmtime(filepath)
                dt_object = datetime.fromtimestamp(mtime)
                # Essayer de lire le premier message pour un aperçu ? (Optionnel, peut ralentir)
                # history_preview = load_log(conv_id)
                # preview = history_preview[0]['content'][:50] + "..." if history_preview else "Vide"
                conversations.append({
                    "id": conv_id,
                    "name": f"Chat du {dt_object.strftime('%d/%m/%Y %H:%M')}", # Nom plus lisible
                    # "preview": preview
                })
            except Exception as e:
                logging.warning(f"Impossible de traiter le fichier log {filename}: {e}")
                conversations.append({"id": conv_id, "name": f"Conversation {conv_id[:8]} (erreur lecture date)"})

        return jsonify({"conversations": conversations})
    except Exception as e:
        logging.error(f"Erreur lors de la lecture de l'historique: {e}", exc_info=True)
        return jsonify({"error": "Impossible de lister l'historique des conversations."}), 500


@app.route('/history/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    """Charge une conversation spécifique depuis l'historique."""
    global current_conversation_history, current_conversation_id, uploaded_file_context

    # Valider l'ID (simple vérification de format UUID si on veut être strict)
    try:
        uuid.UUID(conversation_id, version=4)
    except ValueError:
        logging.warning(f"Tentative de chargement d'historique avec un ID invalide: {conversation_id}")
        return jsonify({"error": "ID de conversation invalide."}), 400

    history = load_log(conversation_id)

    if history:
        # Nettoyer le contexte fichier de l'ancienne conversation active (s'il y en avait une)
        if current_conversation_id and current_conversation_id in uploaded_file_context:
            logging.info(f"Nettoyage du contexte fichier pour l'ancienne conversation {current_conversation_id} lors du chargement de {conversation_id}")
            del uploaded_file_context[current_conversation_id]

        # Le contexte fichier n'est PAS sauvegardé/chargé depuis le log.
        # Donc, on s'assure qu'il n'y a pas de contexte fichier actif pour la conversation chargée.
        if conversation_id in uploaded_file_context:
             logging.info(f"Nettoyage d'un éventuel contexte fichier résiduel pour la conversation chargée {conversation_id}")
             del uploaded_file_context[conversation_id] # Sécurité pour nettoyer un état potentiellement incohérent

        # Mettre à jour l'état global avec la conversation chargée
        current_conversation_id = conversation_id
        current_conversation_history = history
        logging.info(f"Conversation {conversation_id} chargée depuis le log.")

        return jsonify({
            "id": conversation_id,
            "history": history,
            "file_context_status": "Contexte fichier non chargé depuis l'historique. Ré-uploadez le fichier si nécessaire pour cette session."
        })
    else:
        # L'ID était valide mais le fichier log n'a pas été trouvé ou est vide/corrompu
        logging.warning(f"Conversation non trouvée dans les logs pour l'ID: {conversation_id}")
        return jsonify({"error": "Conversation non trouvée ou impossible à charger."}), 404


if __name__ == '__main__':
    # Utiliser host='0.0.0.0' pour rendre accessible sur le réseau local
    # Utiliser host='127.0.0.1' pour un accès strictement local
    # debug=True active le rechargement automatique et le débogueur (NE PAS UTILISER EN PRODUCTION)
    print("--- Démarrage du serveur Flask ---")
    print(f"Modèle Ollama par défaut au démarrage : {current_model}")
    print(f"Accès via : http://127.0.0.1:5000")
    print("---------------------------------")
    app.run(debug=True, host='127.0.0.1', port=5000)